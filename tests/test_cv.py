"""Tests for the CV module — selector, adapter, validator, docx renderer."""

from __future__ import annotations

from html import escape
from pathlib import Path
from types import SimpleNamespace

import pytest

from smartapply.cv import (
    CvAdapter,
    CvBlockSelector,
    CvDocxRenderer,
    CvValidator,
    HtmlApplicationRenderer,
)
from smartapply.cv.education import english_institution_name
from smartapply.cv.motivation_validator import (
    MotivationLetterValidator,
)
from smartapply.llm import (
    AdaptedBullet,
    AdaptedCV,
    AdaptedExperience,
    ApplicationDraft,
    JobAnalysis,
    MockLLMProvider,
    MotivationLetter,
    MotivationLetterRepair,
)
from smartapply.llm.prompts import application_draft, cv_adaptation
from smartapply.pipeline.apply.cv_writer import CvWriterMixin
from smartapply.pipeline.reports import ApplyReport
from smartapply.profile import Degree, get_profile
from smartapply.ranking import MockEmbeddingsProvider


@pytest.fixture(autouse=True)
def _isolated_db(tmp_path: Path, monkeypatch: pytest.MonkeyPatch):
    db_path = tmp_path / "test.db"
    monkeypatch.setenv("DATABASE_URL", f"sqlite:///{db_path}")
    from smartapply.config import get_settings

    get_settings.cache_clear()
    from smartapply.database.session import reset_engine_cache

    reset_engine_cache()
    from smartapply.database.session import init_db

    init_db()
    yield


def _sample_analysis() -> JobAnalysis:
    return JobAnalysis(
        role_type="Data Scientist NLP",
        seniority="mid",
        domain="HealthTech",
        main_tasks=[
            "Build RAG pipelines",
            "Fine-tune transformer models",
            "Productionize ML services",
        ],
        required_skills=["PyTorch", "RAG", "Python"],
        nice_to_have=["Docker"],
        match_reasons=["Strong NLP background", "Multimodal AI experience"],
        risks=["No prior healthtech experience"],
        cv_keywords_to_include=["PyTorch", "RAG", "Whisper", "FAISS"],
    )


def test_selector_batches_all_embedding_inputs_into_one_call() -> None:
    class CountingEmbeddings(MockEmbeddingsProvider):
        def __init__(self):
            self.calls: list[list[str]] = []

        def embed(self, texts: list[str]) -> list[list[float]]:
            self.calls.append(list(texts))
            return super().embed(texts)

    profile = get_profile()
    embeddings = CountingEmbeddings()

    result = CvBlockSelector(embeddings).select(profile, _sample_analysis())

    assert result.experiences
    assert result.projects
    assert len(embeddings.calls) == 1
    assert len(embeddings.calls[0]) == 1 + len(profile.experiences) + len(profile.projects)


def _valid_adapted_cv() -> AdaptedCV:
    return AdaptedCV(
        cv_title="Data Scientist – NLP & Multimodal AI",
        professional_summary=(
            "Data Scientist with 2 years applied R&D in multimodal AI and clinical "
            "digital biomarkers. Strong NLP/RAG and speech pipelines."
        ),
        selected_experiences=[
            AdaptedExperience(
                source_id="exp_emobot_ds_2024",
                bullets=[
                    AdaptedBullet(
                        source_id="blt_emobot_ds_multimodal",
                        text="Built multimodal pipelines reaching 0.67 correlation with clinical scores.",
                    ),
                    AdaptedBullet(
                        source_id="blt_emobot_ds_speech_face",
                        text="Developed speech/NLP and face-recognition pipelines using Whisper and Pyannote.",
                    ),
                ],
            )
        ],
        selected_project_ids=["proj_scifact_rag", "proj_ner_camembert"],
        skills_order=["ml_ai", "data_infra", "stats_signal"],
        warnings=[],
    )


def _empty_skills_cv() -> AdaptedCV:
    return AdaptedCV(
        cv_title="Adapted CV",
        professional_summary="Source-grounded summary.",
        selected_experiences=[],
        selected_project_ids=[],
        skills_profile_id="mixed",
        selected_skills=[],
        skills_order=[],
        warnings=[],
    )


def _analysis_with_required_skills(required: list[str]) -> JobAnalysis:
    return JobAnalysis(
        role_type="Data Specialist",
        seniority="mid",
        domain="Data",
        main_tasks=[],
        required_skills=required,
        nice_to_have=[],
        match_reasons=[],
        risks=[],
        cv_keywords_to_include=[],
    )


def _selected_skill_map(cv: AdaptedCV) -> dict[str, list[str]]:
    return {block.category_id: list(block.skills) for block in cv.selected_skills}


# ---------------- Selector ----------------


# ---------------- Adapter ----------------


def test_offer_skill_matching_prefers_specific_nested_skills() -> None:
    profile = get_profile()
    adapter = CvAdapter(
        profile,
        llm=MockLLMProvider(),
        embeddings=MockEmbeddingsProvider(),
    )

    cases = [
        ("PySpark", "data_infra", ["PySpark"]),
        ("experience with PySpark", "data_infra", ["PySpark"]),
        ("Spark", "data_infra", ["Spark"]),
        ("Mathematical Optimization", "stats_signal", ["Mathematical Optimization"]),
        (
            "advanced mathematical optimization methods",
            "stats_signal",
            ["Mathematical Optimization"],
        ),
        ("Optimization", "stats_signal", ["Optimization"]),
        ("Experience with PyTorch", "ml_ai", ["PyTorch"]),
    ]

    for term, category_id, expected_skills in cases:
        adapted = adapter._ensure_supported_offer_skills(
            _empty_skills_cv(),
            _analysis_with_required_skills([term]),
        )
        selected = _selected_skill_map(adapted)
        assert selected == {category_id: expected_skills}


def test_adapter_can_produce_cv_and_letter_in_one_call() -> None:
    MockLLMProvider.clear()
    cv = _valid_adapted_cv()
    MockLLMProvider.register(
        "application_draft",
        ApplicationDraft(
            cv_title=cv.cv_title,
            professional_summary=cv.professional_summary,
            selected_experiences=cv.selected_experiences,
            selected_project_ids=cv.selected_project_ids,
            skills_order=cv.skills_order,
            warnings=cv.warnings,
            motivation_letter_subject="Application: Data Scientist NLP",
            motivation_letter_body="Hello,\n\nI am writing about this role." + " word" * 180,
        ),
    )
    adapter = CvAdapter(get_profile(), embeddings=MockEmbeddingsProvider())
    adapted, letter, selection = adapter.adapt_application(
        _sample_analysis(), job_title="Data Scientist NLP", job_company="Acme"
    )
    assert adapted.cv_title.startswith("Data Scientist")
    assert letter.subject.startswith("Application")
    assert len(selection.experiences) >= 1


def test_adapter_keeps_project_count_between_two_and_four() -> None:
    adapter = CvAdapter(get_profile(), embeddings=MockEmbeddingsProvider())
    cv = _valid_adapted_cv().model_copy(
        update={
            "selected_project_ids": [
                "proj_scifact_rag",
                "proj_ner_camembert",
                "proj_smartapply",
                "orthogeo3d",
                "proj_gpt2",
            ],
        }
    )

    capped = adapter._enforce_project_count(cv)
    assert capped.selected_project_ids == [
        "proj_scifact_rag",
        "proj_ner_camembert",
        "proj_smartapply",
        "orthogeo3d",
    ]

    short = cv.model_copy(update={"selected_project_ids": ["proj_scifact_rag"]})
    expanded = adapter._enforce_project_count(short)
    assert len(expanded.selected_project_ids) == 2
    assert expanded.selected_project_ids[0] == "proj_scifact_rag"


# ---------------- Validator ----------------


def test_validator_accepts_clean_cv() -> None:
    v = CvValidator(get_profile())
    result = v.validate(_valid_adapted_cv())
    assert result.ok
    assert not result.errors


def test_validator_rejects_unknown_bullet_id() -> None:
    cv = _valid_adapted_cv()
    cv.selected_experiences[0].bullets[0] = AdaptedBullet(
        source_id="blt_does_not_exist",
        text="Did something cool.",
    )
    v = CvValidator(get_profile())
    result = v.validate(cv)
    assert not result.ok
    assert any("unknown_bullet_id" in e for e in result.errors)


def test_motivation_letter_validator_flags_short_and_unsupported_terms() -> None:
    analysis = _sample_analysis().model_copy(
        update={"required_skills": ["Kubernetes"], "cv_keywords_to_include": []}
    )
    letter = MotivationLetter(
        subject="Application: Data Scientist",
        body="My experience includes Kubernetes.",
    )
    result = MotivationLetterValidator(get_profile()).validate(
        letter,
        cv=_valid_adapted_cv(),
        analysis=analysis,
    )
    assert "letter_too_short:4" in result.warnings
    assert "unsupported_term_in_letter:Kubernetes" in result.warnings


def test_cv_title_alignment_compares_role_families_not_offer_keywords() -> None:
    analysis = _sample_analysis().model_copy(
        update={
            "role_type": "Machine Learning Engineer",
            "main_tasks": ["Build RAG pipelines"],
            "required_skills": ["Python", "LangChain", "Kubernetes"],
            "cv_keywords_to_include": ["RAG", "vector search"],
        }
    )
    cv = _valid_adapted_cv().model_copy(
        update={
            "cv_title": "AI Engineer - Applied Intelligence",
            "professional_summary": "Machine learning profile for grounded retrieval workflows.",
        }
    )
    report = ApplyReport(job_id=42, application_id=None)

    CvWriterMixin()._validate_cv_offer_alignment(
        cv,
        analysis,
        "Machine Learning Engineer",
        report,
    )

    assert "cv_title_not_offer_anchored" not in report.validation_warnings


def test_cv_title_alignment_rejects_known_incompatible_family() -> None:
    analysis = _sample_analysis().model_copy(
        update={
            "role_type": "Data Engineer",
            "main_tasks": ["Build data pipelines"],
            "required_skills": ["Kubernetes"],
            "cv_keywords_to_include": [],
        }
    )
    cv = _valid_adapted_cv().model_copy(
        update={
            "cv_title": "Data Scientist - Kubernetes",
            "professional_summary": "Data engineering profile focused on reliable pipelines.",
        }
    )
    report = ApplyReport(job_id=42, application_id=None)

    CvWriterMixin()._validate_cv_offer_alignment(
        cv,
        analysis,
        "Data Engineer",
        report,
    )

    assert "cv_title_not_offer_anchored" in report.validation_warnings


def test_cv_title_alignment_ignores_ambiguous_title_family() -> None:
    analysis = _sample_analysis().model_copy(update={"role_type": "Data Engineer"})
    cv = _valid_adapted_cv().model_copy(update={"cv_title": "Applied Intelligence Specialist"})
    report = ApplyReport(job_id=42, application_id=None)

    CvWriterMixin()._validate_cv_offer_alignment(
        cv,
        analysis,
        "Data Engineer",
        report,
    )

    assert "cv_title_not_offer_anchored" not in report.validation_warnings


class _RepairRecordingProvider(MockLLMProvider):
    def __init__(self, registry=None):  # noqa: ANN001
        super().__init__(registry)
        self.calls: list[dict] = []

    def complete_json(self, **kwargs):  # noqa: ANN003, ANN201
        self.calls.append(kwargs)
        return super().complete_json(**kwargs)


class _RepairHarness(CvWriterMixin):
    def __init__(self, provider: MockLLMProvider):
        self.profile = get_profile()
        self.llm = provider
        self.letter_validator = MotivationLetterValidator(self.profile)


def _repair_valid_body() -> str:
    refrain = (
        "I value careful analysis, clear communication, reliable delivery, and "
        "thoughtful collaboration."
    )
    openings = (
        "The SciFact RAG Verifier project strengthened my approach to grounded "
        "retrieval, evidence review, and dependable evaluation.",
        "This selected work reflects a practical way of connecting technical "
        "decisions with clear product priorities and team needs.",
        "For this role, I would bring the same structured approach, attention "
        "to quality, and focus on useful outcomes.",
    )
    return "\n\n".join(f"{opening} {' '.join([refrain] * 6)}" for opening in openings)


def _repair_too_long_body() -> str:
    return f"{_repair_valid_body()} {' focused delivery' * 80}"


def _repair_once(
    harness: _RepairHarness,
    letter: MotivationLetter,
) -> MotivationLetter:
    return harness._repair_letter_once(
        letter,
        _valid_adapted_cv(),
        _sample_analysis(),
        job_title="Data Scientist",
        job_company="Acme",
        language="en",
        job_id=42,
    )


def test_single_letter_defect_is_repaired_once_with_cheap_model() -> None:
    provider = _RepairRecordingProvider(
        {
            "motivation_letter_repair": MotivationLetterRepair(body=_repair_valid_body()),
        }
    )
    harness = _RepairHarness(provider)
    original = MotivationLetter(
        subject="Application - Data Scientist",
        body=_repair_too_long_body(),
    )

    repaired = _repair_once(harness, original)

    assert repaired.subject == original.subject
    assert repaired.body == _repair_valid_body()
    assert len(provider.calls) == 1
    call = provider.calls[0]
    assert call["purpose"] == "motivation_letter_repair"
    assert call["model"] == provider.cheap_model
    assert call["temperature"] == 0.1
    assert call["schema"] is MotivationLetterRepair
    assert "letter_too_long" in call["user"]
    result = harness.letter_validator.validate(
        repaired,
        cv=_valid_adapted_cv(),
        analysis=_sample_analysis(),
    )
    assert result.errors == []
    assert result.warnings == []


def test_multiple_letter_defects_do_not_trigger_repair() -> None:
    provider = _RepairRecordingProvider()
    harness = _RepairHarness(provider)
    original = MotivationLetter(subject="Application", body="Brief note.")

    repaired = _repair_once(harness, original)

    assert repaired == original
    assert provider.calls == []


def test_unsuccessful_letter_repair_is_rejected_after_single_attempt() -> None:
    provider = _RepairRecordingProvider(
        {
            "motivation_letter_repair": MotivationLetterRepair(body="Still too brief."),
        }
    )
    harness = _RepairHarness(provider)
    original = MotivationLetter(subject="Application", body=_repair_too_long_body())

    repaired = _repair_once(harness, original)

    assert repaired == original
    assert len(provider.calls) == 1


def test_letter_repair_provider_failure_keeps_original() -> None:
    provider = _RepairRecordingProvider()
    harness = _RepairHarness(provider)
    original = MotivationLetter(subject="Application", body=_repair_too_long_body())

    repaired = _repair_once(harness, original)

    assert repaired == original
    assert len(provider.calls) == 1


# ---------------- DOCX renderer ----------------


def test_docx_renderer_produces_valid_file(tmp_path: Path) -> None:
    profile = get_profile()
    cv = _valid_adapted_cv()
    out = tmp_path / "cv.docx"
    renderer = CvDocxRenderer(profile)
    renderer.save(cv, out)
    assert out.exists()
    # File starts with the ZIP magic bytes (DOCX is a ZIP)
    assert out.read_bytes()[:2] == b"PK"
    # Re-open to confirm validity
    from docx import Document

    doc = Document(str(out))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Lachtar Nour" in full_text
    assert "Data Scientist" in full_text


def test_html_renderer_matches_reference_cv_architecture() -> None:
    profile = get_profile()
    html = HtmlApplicationRenderer(profile).render_cv_html(_valid_adapted_cv())
    assert "Professional Experience" in html
    assert "Education" in html
    assert "Skills" in html
    assert "Projects" in html
    assert "Languages" in html
    assert "Certificates" in html
    assert "section-title" in html
    assert "header-grid" in html
    assert "identity-line" in html
    assert "grid-template-columns: max-content 1px minmax(0, 70mm)" in html
    assert "max-width: 70mm" in html
    assert "overflow-wrap: anywhere" in html
    assert "white-space: normal" in html
    assert "--header-gap: 20px" in html
    assert "--summary-experience-gap: 6px" in html
    assert "calc(var(--header-gap) + var(--summary-experience-gap))" in html
    assert "justify-content: var(--section-justify)" in html
    assert "--section-content-indent: 3.5mm" in html
    assert "grid-template-columns: 34mm 1fr" in html
    assert "skill-row" in html
    assert "skill-pill" in html
    assert "language-row" in html
    assert "border-left" not in html
    assert "header::before" not in html
    assert "\\2197" not in html
    assert 'class="project-link"' in html
    assert 'aria-label="Project repository"' in html
    assert "Completed" not in html


def test_letter_renderer_uses_adapted_headline_and_fallback() -> None:
    profile = get_profile()
    renderer = HtmlApplicationRenderer(profile)
    letter = MotivationLetter(
        subject="Candidature - Data Scientist",
        body="Premier paragraphe.\n\nDeuxième paragraphe.\n\nTroisième paragraphe.",
    )

    html = renderer.render_letter_html(
        letter=letter,
        job_title="Data Scientist",
        job_company="Acme",
        letter_headline="Data Scientist - NLP & Multimodal AI",
    )

    assert '<p class="headline">Data Scientist - NLP &amp; Multimodal AI</p>' in html
    profile_headline = escape(profile.identity.title)
    assert f'<p class="headline">{profile_headline}</p>' not in html

    fallback_html = renderer.render_letter_html(
        letter=letter,
        job_title="Data Scientist",
        job_company="Acme",
    )

    assert f'<p class="headline">{profile_headline}</p>' in fallback_html


def test_application_renderer_passes_cv_title_to_letter_renderer(tmp_path: Path) -> None:
    from smartapply.pipeline.application_renderer import ApplicationDocumentRenderer
    from smartapply.pipeline.reports import ApplyReport

    profile = get_profile()
    adapted = _valid_adapted_cv()
    original_cv_title = adapted.cv_title
    letter = MotivationLetter(
        subject="Candidature - Data Scientist",
        body="Premier paragraphe.\n\nDeuxième paragraphe.\n\nTroisième paragraphe.",
    )
    letter_calls: list[tuple[str, str | None]] = []
    cv_titles_seen: list[str] = []

    class FakeDocxRenderer:
        def save(self, adapted_cv, path):  # noqa: ANN001
            Path(path).parent.mkdir(parents=True, exist_ok=True)
            Path(path).write_text("docx", encoding="utf-8")

    class FakeHtmlRenderer:
        def save_cv_html(self, adapted_cv, path):  # noqa: ANN001
            cv_titles_seen.append(adapted_cv.cv_title)
            Path(path).parent.mkdir(parents=True, exist_ok=True)
            Path(path).write_text("cv html", encoding="utf-8")

        def save_cv_pdf(self, adapted_cv, path):  # noqa: ANN001
            cv_titles_seen.append(adapted_cv.cv_title)
            Path(path).parent.mkdir(parents=True, exist_ok=True)
            Path(path).write_text("cv pdf", encoding="utf-8")

        def save_letter_html(self, **kwargs):  # noqa: ANN001
            letter_calls.append(("html", kwargs.get("letter_headline")))
            path = Path(kwargs["path"])
            path.parent.mkdir(parents=True, exist_ok=True)
            path.write_text("letter html", encoding="utf-8")

        def save_letter_pdf(self, **kwargs):  # noqa: ANN001
            letter_calls.append(("pdf", kwargs.get("letter_headline")))
            path = Path(kwargs["path"])
            path.parent.mkdir(parents=True, exist_ok=True)
            path.write_text("letter pdf", encoding="utf-8")

    renderer = ApplicationDocumentRenderer.__new__(ApplicationDocumentRenderer)
    renderer.profile = profile
    renderer.docx = FakeDocxRenderer()
    renderer.html = FakeHtmlRenderer()
    renderer.settings = SimpleNamespace(output_dir=tmp_path)
    report = ApplyReport(job_id=1, application_id=123)

    renderer.render_all(
        report=report,
        adapted=adapted,
        letter=letter,
        job_title="Data Scientist",
        job_company="Acme",
        language="fr",
    )

    assert letter_calls == [
        ("html", original_cv_title),
        ("pdf", original_cv_title),
    ]
    assert cv_titles_seen == [original_cv_title, original_cv_title]
    assert adapted.cv_title == original_cv_title


def test_cv_renderers_translate_french_university_label(tmp_path: Path) -> None:
    assert english_institution_name("Sorbonne Université") == "Sorbonne University"

    profile = get_profile().model_copy(
        update={
            "education": [
                Degree(
                    id="edu_paris_cite",
                    title="Master's Degree",
                    field="Data Science",
                    institution="Université Paris Cité",
                    start_date="09/2022",
                    end_date="09/2024",
                    start_year=2022,
                    end_year=2024,
                )
            ]
        }
    )

    html = HtmlApplicationRenderer(profile).render_cv_html(_valid_adapted_cv())
    assert "Paris Cité University" in html
    assert "Université" not in html
    assert "09/2022 &ndash; 09/2024" in html

    out = tmp_path / "cv.docx"
    CvDocxRenderer(profile).save(_valid_adapted_cv(), out)

    from docx import Document

    doc = Document(str(out))
    full_text = "\n".join(
        p.text
        for table in doc.tables
        for row in table.rows
        for cell in row.cells
        for p in cell.paragraphs
    )
    assert "Paris Cité University" in full_text
    assert "Université" not in full_text
    assert "09/2022 – 09/2024" in full_text


def test_application_draft_prompt_separates_matching_keywords_from_display_skills() -> None:
    profile = get_profile()
    analysis = _sample_analysis().model_copy(
        update={
            "required_skills": [
                "Python",
                "Machine learning",
                "Computer vision",
                "MLOps",
                "Kubernetes",
                "Experience with PyTorch",
            ],
            "cv_keywords_to_include": ["PyTorch", "GCP"],
            "company_context": (
                "Acme Health builds patient-facing clinical AI products for care teams."
            ),
            "offer_interest_points": [
                "Work on RAG pipelines for medical knowledge access",
                "Collaborate with product and clinical teams",
            ],
        }
    )
    app_prompt = application_draft.build_user_prompt(
        profile=profile,
        analysis=analysis,
        job_title="Clinical AI Engineer",
        job_company="Acme Health",
        selected_experiences=profile.experiences[:1],
        selected_projects=profile.projects[:2],
        language="en",
    )
    cv_prompt = cv_adaptation.build_user_prompt(
        profile=profile,
        analysis=analysis,
        job_title="Clinical AI Engineer",
        job_company="Acme Health",
        selected_experiences=profile.experiences[:1],
        selected_projects=profile.projects[:2],
    )
    combined = "\n".join(
        [
            application_draft.SYSTEM,
            app_prompt,
            cv_adaptation.SYSTEM,
            cv_prompt,
        ]
    )

    assert "allowed_skills_by_category" in app_prompt
    assert "core_skills_by_category" in app_prompt
    assert "matching_keywords_for_profile_selection_not_display" in app_prompt
    assert "genai" in app_prompt
    assert "do not copy those keywords into selected_skills" in app_prompt
    assert "unsupported_offer_terms_not_to_claim" in app_prompt
    assert "- MLOps" in app_prompt
    assert "- Kubernetes" in app_prompt
    assert "- GCP" in app_prompt
    assert "- Machine learning" not in app_prompt
    assert "- Computer vision" not in app_prompt
    assert "Offer/company anchors for letter:" in app_prompt
    assert "Acme Health builds patient-facing clinical AI products for care teams." in app_prompt
    assert "Work on RAG pipelines for medical knowledge access" in app_prompt
    assert "sound like a real candidate" in app_prompt
    assert "why this company and this role specifically make sense" in app_prompt
    assert "=== MOTIVATION LETTER ===" in app_prompt
    assert "French or English: 220-300 words" in app_prompt
    assert "clear recruiter-facing language" in app_prompt
    assert "why it matters for the role" in application_draft.SYSTEM
    assert "Do not reproduce or closely paraphrase CV bullets" in app_prompt
    assert (
        "For professional experiences, use past or clearly neutral wording only"
        in application_draft.SYSTEM
    )
    assert "never describe them as current" in application_draft.SYSTEM
    assert (
        "Projects marked current or ongoing may still be described as current work"
        in application_draft.SYSTEM
    )
    assert "During my experience at Emobot" in app_prompt
    assert "At Emobot, I developed" in app_prompt
    assert "Chez Emobot, j’ai développé" in app_prompt
    assert "Lors de mon expérience chez Emobot" in app_prompt
    assert "my current role" in app_prompt
    assert "currently at Emobot" in app_prompt
    assert "dans mon poste actuel" in app_prompt
    assert "actuellement chez Emobot" in app_prompt
    assert "Projects marked current or ongoing may still be described as current work" in app_prompt
    assert "In my role at" not in app_prompt
    assert "one concrete priority of the role or company" in application_draft.SYSTEM
    assert "Avoid generic closing language" in application_draft.SYSTEM
    assert "Only claim skills, tools, domains, or experience" in app_prompt
    assert "Clear link between company needs, candidate evidence" in app_prompt
    assert "selected_project_ids should contain 2 to 4 projects" in app_prompt
    assert "Do not add weak filler projects just to reach 4" in app_prompt
    assert "cv_title MUST read like a polished, real CV headline" in application_draft.SYSTEM
    assert "not a skills list" in application_draft.SYSTEM
    assert "not tools or stack keywords" in app_prompt
    assert "opening role consistent with cv_title" in app_prompt
    assert 'under an "AI Engineer" title' in app_prompt
    assert (
        "CV language: English only. Keep cv_title, professional_summary and CV bullets in English."
        in app_prompt
    )
    assert "CV title: write a stable, polished professional headline" in app_prompt
    assert (
        "Experience output: include every provided experience and every bullet id exactly once"
        in app_prompt
    )
    assert "write a motivation letter" in application_draft.SYSTEM
    assert "Required skills (for selected_skills, professional_summary and bullets" not in combined
    assert (
        "Required skills (from offer; for selected_skills, professional_summary and bullets"
        not in combined
    )
    assert "Keywords to surface in professional_summary" not in combined
    assert "Include supported technical anchors" not in combined
    assert (
        "Did you move offer-specific tasks, tools and keywords to professional_summary"
        not in combined
    )
    assert "not to build a keyword list" in combined
    assert "Do not use sequences of tools or technologies for keyword coverage" in combined
    assert (
        "A named model family, method, project or implementation technology may "
        "appear only when it is integral to a specific source-grounded contribution"
    ) in combined
    assert "Do not turn the summary into a stack inventory or keyword list" in combined
    assert (
        "named technologies, methods or projects must be tied to concrete source-grounded work"
        in combined
    )
    assert (
        "Use allowed_skills to verify that any named method, model family, "
        "project or implementation technology is supported; do not enumerate "
        "skills as standalone content"
    ) in combined
    assert (
        "professional_summary MUST be recruiter-facing, concrete, adapted to the role" in combined
    )
    assert (
        "reflect at least one supported role-specific work area, workflow or contribution"
        in combined
    )
    assert "selected_skills is the exact Skills section to display" in combined
    assert (
        "Include every required or strongly requested offer skill that exists in allowed_skills"
        in combined
    )
    assert (
        "CV title: write a stable, polished professional headline under the candidate's name"
    ) in combined
    assert (
        "Keywords to interpret the role focus and surface in selected_skills, "
        "source-grounded bullets and the motivation letter when supported"
    ) in app_prompt
    assert "exact skills" not in combined
