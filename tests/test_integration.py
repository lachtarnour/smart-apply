"""End-to-end integration test on a realistic sample of jobs.

Drives the entire pipeline with mocked LLM — no API keys, no network.
Verifies anti-hallucination, persistence, and artifact files.
"""

from __future__ import annotations

from pathlib import Path

import pytest

from smartapply.llm import (
    AdaptedBullet,
    AdaptedCV,
    AdaptedExperience,
    ApplicationDraft,
    JobAnalysis,
    MockLLMProvider,
)
from smartapply.ranking import MockEmbeddingsProvider


@pytest.fixture(autouse=True)
def _isolated_env(tmp_path: Path, monkeypatch: pytest.MonkeyPatch):
    db_path = tmp_path / "smartapply.db"
    output_dir = tmp_path / "output"
    monkeypatch.setenv("DATABASE_URL", f"sqlite:///{db_path}")
    monkeypatch.setenv("OUTPUT_DIR", str(output_dir))

    from smartapply.config import get_settings

    get_settings.cache_clear()
    from smartapply.database.session import reset_engine_cache

    reset_engine_cache()
    from smartapply.database.session import init_db

    init_db()
    yield


SAMPLE_JOBS = [
    {
        "title": "Data Scientist NLP",
        "company": "Doctolib",
        "location": "Paris, France",
        "contract_type": "CDI",
        "remote_policy": "hybrid",
        "application_url": "https://careers.doctolib.com/jobs/42",
        "description": (
            "Vos missions\n"
            "Concevoir des pipelines RAG (BM25, FAISS, reranking) sur des données "
            "médicales. Fine-tuner des modèles Hugging Face. Déployer en production "
            "via Docker et AWS. Travailler avec PyTorch et Whisper pour la voix.\n\n"
            "Profil recherché\n"
            "2-4 ans d'expérience en ML. Maîtrise Python, PyTorch. NLP solide. "
            "Connaissance des outils LLM (Hugging Face, LangChain). Anglais courant."
        ),
    },
    {
        "title": "Senior ML Engineer",
        "company": "Mistral AI",
        "location": "Paris, France",
        "contract_type": "CDI",
        "remote_policy": "onsite",
        "application_url": "https://jobs.mistral.ai/ml-engineer",
        "description": (
            "Build distributed training infrastructure for foundation models. "
            "5+ years of experience required. PyTorch, JAX, CUDA expert. "
            "PhD preferred."
        ),
    },
    {
        "title": "Data Analyst BI",
        "company": "Carrefour",
        "location": "Massy, France",
        "contract_type": "CDI",
        "remote_policy": "hybrid",
        "application_url": "https://careers.carrefour.com/analyst",
        "description": (
            "Création de dashboards Power BI. Reporting financier mensuel. "
            "SQL avancé requis. Pas de développement Python."
        ),
    },
    {
        "title": "Sales Director Europe",
        "company": "Salesforce",
        "location": "Paris, France",
        "contract_type": "CDI",
        "remote_policy": "hybrid",
        "application_url": "https://salesforce.com/careers",
        "description": (
            "Lead the European B2B sales organization. 10+ years of experience "
            "managing enterprise sales teams. Hit aggressive growth targets."
        ),
    },
    {
        "title": "AI Research Engineer",
        "company": "Hugging Face",
        "location": "Remote (EU)",
        "contract_type": "CDI",
        "remote_policy": "remote",
        "application_url": "https://huggingface.co/jobs/ai-researcher",
        "description": (
            "Research engineer to push the state of the art in multimodal AI. "
            "Strong publications and engineering skills. PyTorch, Transformers, "
            "Diffusion models. 2-5 years experience."
        ),
    },
]


def _register_llm_fixtures() -> None:
    MockLLMProvider.clear()
    MockLLMProvider.register(
        "job_analysis",
        JobAnalysis(
            role_type="Data Scientist NLP",
            seniority="mid",
            domain="HealthTech",
            main_tasks=[
                "Concevoir des pipelines RAG",
                "Fine-tuner des modèles transformers",
                "Déployer en production",
            ],
            required_skills=["Python", "PyTorch", "RAG", "FAISS"],
            nice_to_have=["AWS", "Docker"],
            match_reasons=[
                "Forte expérience NLP/RAG",
                "Projet Evidence RAG",
                "Pipelines speech via Whisper",
            ],
            risks=["Pas d'expérience prod cloud à grande échelle"],
            cv_keywords_to_include=[
                "PyTorch",
                "Hugging Face",
                "FAISS",
                "BM25",
                "Whisper",
                "RAG",
            ],
        ),
    )
    MockLLMProvider.register(
        "cv_adaptation",
        AdaptedCV(
            cv_title="Data Scientist – NLP & Multimodal AI",
            professional_summary=(
                "Data Scientist with 2 years applied R&D in NLP, multimodal AI and "
                "clinical digital biomarkers. Built RAG pipelines, speech/NLP "
                "stacks (Whisper, Pyannote) and multimodal models reaching 0.67 "
                "correlation with clinical scores."
            ),
            selected_experiences=[
                AdaptedExperience(
                    source_id="exp_aurore_ds_2024",
                    bullets=[
                        AdaptedBullet(
                            source_id="blt_aurore_ds_multimodal",
                            text=(
                                "Built multimodal digital biomarker pipelines from facial, "
                                "mobility and smartphone data, reaching 0.67 correlation "
                                "with validated clinical scores."
                            ),
                        ),
                        AdaptedBullet(
                            source_id="blt_aurore_ds_speech_face",
                            text=(
                                "Developed speech/NLP and face-recognition pipelines using "
                                "Whisper, Pyannote, RetinaFace, FaceNet and Flask APIs."
                            ),
                        ),
                        AdaptedBullet(
                            source_id="blt_aurore_ds_patent",
                            text=(
                                "Contributed to a patent-pending AI monitoring system and "
                                "clinical preprint on passive mood markers."
                            ),
                        ),
                    ],
                ),
                AdaptedExperience(
                    source_id="exp_aurore_intern_2023",
                    bullets=[
                        AdaptedBullet(
                            source_id="blt_aurore_intern_anomaly",
                            text=(
                                "Built an anomaly detection pipeline for identifying "
                                "behavioral disruptions in mood tracking data."
                            ),
                        ),
                    ],
                ),
            ],
            selected_project_ids=[
                "proj_evidence_rag",
                "proj_ner_camembert",
                "proj_gpt2",
            ],
            skills_order=["ml_ai", "data_infra", "stats_signal"],
            warnings=[],
        ),
    )
    MockLLMProvider.register(
        "application_draft",
        ApplicationDraft(
            cv_title="Data Scientist – NLP & Multimodal AI",
            professional_summary=(
                "Data Scientist with 2 years applied R&D in NLP, multimodal AI and "
                "clinical digital biomarkers. Built RAG pipelines, speech/NLP "
                "stacks (Whisper, Pyannote) and multimodal models reaching 0.67 "
                "correlation with clinical scores."
            ),
            selected_experiences=[
                AdaptedExperience(
                    source_id="exp_aurore_ds_2024",
                    bullets=[
                        AdaptedBullet(
                            source_id="blt_aurore_ds_multimodal",
                            text=(
                                "Built multimodal digital biomarker pipelines from facial, "
                                "mobility and smartphone data, reaching 0.67 correlation "
                                "with validated clinical scores."
                            ),
                        ),
                        AdaptedBullet(
                            source_id="blt_aurore_ds_speech_face",
                            text=(
                                "Developed speech/NLP and face-recognition pipelines using "
                                "Whisper, Pyannote, RetinaFace, FaceNet and Flask APIs."
                            ),
                        ),
                        AdaptedBullet(
                            source_id="blt_aurore_ds_patent",
                            text=(
                                "Contributed to a patent-pending AI monitoring system and "
                                "clinical preprint on passive mood markers."
                            ),
                        ),
                    ],
                ),
                AdaptedExperience(
                    source_id="exp_aurore_intern_2023",
                    bullets=[
                        AdaptedBullet(
                            source_id="blt_aurore_intern_anomaly",
                            text=(
                                "Built an anomaly detection pipeline for identifying "
                                "behavioral disruptions in mood tracking data."
                            ),
                        ),
                    ],
                ),
            ],
            selected_project_ids=[
                "proj_evidence_rag",
                "proj_ner_camembert",
                "proj_gpt2",
            ],
            skills_order=["ml_ai", "data_infra", "stats_signal"],
            warnings=[],
            motivation_letter_subject="Candidature - Data Scientist NLP - Camille Martin",
            motivation_letter_body=(
                "Bonjour,\n\n"
                "Je vous adresse ma candidature pour le poste de Data Scientist NLP. "
                "Mes deux années chez Aurore Labs m'ont permis de construire des pipelines "
                "NLP et speech avec Whisper et Pyannote, ainsi que des biomarqueurs "
                "cliniques multimodaux atteignant 0.67 de corrélation avec des scores "
                "validés. Le projet Evidence RAG, fondé sur BM25, FAISS, reranking et "
                "génération de réponses sourcées, rejoint directement les missions de "
                "RAG, de fine-tuning et de déploiement que vous décrivez. Ce parcours "
                "combine expérimentation, évaluation et intégration logicielle sur des "
                "données complexes. Je serais ravi d'échanger sur la manière dont ce "
                "profil peut contribuer à vos sujets NLP et IA appliquée.\n\n"
                "Cordialement,\n"
                "Camille Martin"
            ),
        ),
    )


def test_full_pipeline_on_realistic_sample(tmp_path: Path) -> None:
    """Ingest 5 jobs → process → apply to top 2. Verify everything."""
    _register_llm_fixtures()

    from smartapply.database import session_scope
    from smartapply.database.repository import top_jobs_by_score
    from smartapply.pipeline import Pipeline

    p = Pipeline(embeddings=MockEmbeddingsProvider(), llm=MockLLMProvider())

    # ---- Step 1: Ingest the 5 sample jobs ----
    for job_data in SAMPLE_JOBS:
        p.ingest_text(
            text=job_data["description"],
            title=job_data["title"],
            company=job_data["company"],
            location=job_data["location"],
            application_url=job_data["application_url"],
        )

    # ---- Step 2: Process ----
    process_report = p.process_pending()
    assert process_report.total == 5
    # Sales Director should be filtered out by the rules engine.
    # Data Scientist NLP + AI Research Engineer must be analyzed.
    assert process_report.kept_after_filter >= 2
    assert process_report.kept_after_filter <= 4
    assert process_report.analyzed >= 2

    # ---- Step 3: Apply to top-scoring jobs ----
    with session_scope() as s:
        top = list(top_jobs_by_score(s, 2))
        top_ids = [j.id for j in top]
        top_companies = [j.company for j in top]

    assert top_ids
    # The Sales Director should NOT be in the top
    assert "Salesforce" not in top_companies
    # The BI role should NOT be in the top
    assert "Carrefour" not in top_companies

    reports = []
    for jid in top_ids:
        report = p.apply_to(jid)
        reports.append(report)
        assert report.application_id is not None
        assert report.docx_path and Path(report.docx_path).exists()
        assert report.cv_pdf_path and Path(report.cv_pdf_path).exists()
        assert report.letter_pdf_path and Path(report.letter_pdf_path).exists()
        assert Path(report.docx_path).parent.name == str(report.application_id)
        # Anti-hallucination: no validation errors after auto-fix
        assert not report.validation_errors

    # ---- Step 4: Verify generated CV contains the fixture's profile facts ----
    from docx import Document

    docx_path = reports[0].docx_path
    doc = Document(docx_path)
    text = "\n".join(p.text for p in doc.paragraphs)
    # Extract text from tables too
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text += "\n" + cell.text

    # Fixture facts from the profile must appear
    assert "Camille Martin" in text
    assert "Aurore Labs" in text
    assert "0.67" in text  # the validated quantified metric
    assert "Whisper" in text  # real tech from profile
    # Allowed skills should appear
    assert "PyTorch" in text

    # ---- Step 5: Application persisted with its complete document dossier ----
    from smartapply.database.models import Application

    with session_scope() as s:
        app = s.get(Application, reports[0].application_id)
        assert app is not None
        assert app.cv_json  # full structured CV is persisted
        assert app.status == "ready_for_form_submission"
        assert app.cv_pdf_path and Path(app.cv_pdf_path).exists()
        assert len(app.documents) >= 7  # CV, letter, HTML/PDF sources and CV JSON
