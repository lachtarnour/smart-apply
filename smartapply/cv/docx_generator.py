"""DOCX renderer matching the user's existing CV template.

Layout:
- Header: large name, italic summary, contact line.
- For each section (Professional Experience, Education, Skills, Projects,
  Languages, Certificates): blue underline title, then a 2-column table
  (left: dates+location, right: title + bullets).
"""

from __future__ import annotations

from collections.abc import Callable
from pathlib import Path

from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from smartapply.cv.education import EducationDisplay, education_entries_for_english
from smartapply.cv.links import split_bullet_with_links
from smartapply.cv.skill_profile import infer_skill_profile_id
from smartapply.llm import AdaptedCV
from smartapply.profile import (
    BulletLink,
    Profile,
    Project,
    TemplateStyle,
)

# -------------------- XML helpers --------------------


def _add_bottom_border(paragraph, color_hex: str, size: int = 8) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _remove_table_borders(table) -> None:
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "nil")
        tblBorders.append(el)
    tblPr.append(tblBorders)


def _add_hyperlink_run(paragraph, text: str, url: str, font_size: Pt) -> None:
    """Append a hyperlinked run to ``paragraph`` styled like a blue underline link."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(font_size.pt * 2)))
    rPr.append(sz)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "1F4FD1")
    rPr.append(color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    new_run.append(rPr)

    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


# -------------------- Renderer --------------------


class CvDocxRenderer:
    """Render an ``AdaptedCV`` (or full profile) to a DOCX file."""

    def __init__(self, profile: Profile, style: TemplateStyle | None = None):
        self.profile = profile
        self.style = style or profile.template_style

    # ---- public API ----

    def render(self, adapted: AdaptedCV) -> DocxDocument:
        doc = Document()
        self._set_margins(doc)
        self._write_header(doc, adapted)
        self._write_section(
            doc, "Professional Experience", lambda: self._write_experiences(doc, adapted)
        )
        if self.profile.education:
            self._write_section(doc, "Education", lambda: self._write_education(doc))
        self._write_section(doc, "Skills", lambda: self._write_skills(doc, adapted))
        if adapted.selected_project_ids and self.profile.projects:
            self._write_section(doc, "Projects", lambda: self._write_projects(doc, adapted))
        if self.profile.languages:
            self._write_section(doc, "Languages", lambda: self._write_languages(doc))
        if self.profile.certificates:
            self._write_section(doc, "Certificates", lambda: self._write_certificates(doc))
        return doc

    def save(self, adapted: AdaptedCV, path: str | Path) -> Path:
        doc = self.render(adapted)
        out = Path(path)
        out.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(out))
        return out

    # ---- internals ----

    def _set_margins(self, doc: DocxDocument) -> None:
        for section in doc.sections:
            section.top_margin = Cm(1.5)
            section.bottom_margin = Cm(1.5)
            section.left_margin = Cm(1.8)
            section.right_margin = Cm(1.8)

    def _write_header(self, doc: DocxDocument, adapted: AdaptedCV) -> None:
        p_name = doc.add_paragraph()
        run = p_name.add_run(self.profile.identity.full_name)
        run.bold = True
        run.font.size = Pt(self.style.title_font_size)
        run.font.name = self.style.font_family
        run.font.color.rgb = RGBColor.from_string(self.style.text_color_hex)

        if adapted.cv_title:
            p_title = doc.add_paragraph()
            r = p_title.add_run(adapted.cv_title)
            r.italic = True
            r.font.size = Pt(self.style.body_font_size + 1)
            r.font.color.rgb = RGBColor.from_string(self.style.muted_color_hex)

        if adapted.professional_summary:
            p_summary = doc.add_paragraph()
            r = p_summary.add_run(adapted.professional_summary)
            r.italic = True
            r.font.size = Pt(self.style.body_font_size)
            r.font.color.rgb = RGBColor.from_string(self.style.text_color_hex)

        contact_bits = [
            self.profile.identity.location,
            self.profile.identity.phone,
            self.profile.identity.email,
            self.profile.identity.github,
            self.profile.identity.linkedin,
            self.profile.identity.website,
        ]
        contact = "  |  ".join(str(b) for b in contact_bits if b)
        p_contact = doc.add_paragraph()
        r = p_contact.add_run(contact)
        r.font.size = Pt(self.style.body_font_size)
        r.font.color.rgb = RGBColor.from_string(self.style.muted_color_hex)
        p_contact.paragraph_format.space_after = Pt(6)

    def _write_section(self, doc: DocxDocument, title: str, body: Callable[[], None]) -> None:
        doc.add_paragraph()  # spacer
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(self.style.section_font_size)
        run.font.color.rgb = RGBColor.from_string(self.style.primary_color_hex)
        run.font.name = self.style.font_family
        _add_bottom_border(p, self.style.primary_color_hex)
        body()

    def _two_col_row(self, doc: DocxDocument, left: str, right_writer: Callable[[object], None]) -> None:
        table = doc.add_table(rows=1, cols=2)
        _remove_table_borders(table)
        table.autofit = False
        left_w = Cm(self.style.left_column_width_cm)
        right_w = Cm(16.5 - self.style.left_column_width_cm)
        table.columns[0].width = left_w
        table.columns[1].width = right_w
        cell_left, cell_right = table.rows[0].cells
        cell_left.width = left_w
        cell_right.width = right_w

        p_left = cell_left.paragraphs[0]
        r = p_left.add_run(left)
        r.font.size = Pt(self.style.body_font_size)
        r.font.color.rgb = RGBColor.from_string(self.style.muted_color_hex)

        right_writer(cell_right)

    # ---- sections content ----

    def _write_experiences(self, doc: DocxDocument, adapted: AdaptedCV) -> None:
        exp_by_id = {e.id: e for e in self.profile.experiences}
        for exp_adapted in adapted.selected_experiences:
            src = exp_by_id.get(exp_adapted.source_id)
            if src is None:
                continue
            left = f"{src.start_date} – {src.end_date}\n{src.location}"
            source_links_by_id = {b.id: b.links for b in src.bullets}

            def _writer(cell, src=src, exp_adapted=exp_adapted, links_by_id=source_links_by_id):
                p_title = cell.paragraphs[0]
                r1 = p_title.add_run(f"{src.title}, ")
                r1.bold = True
                r1.font.size = Pt(self.style.body_font_size + 0.5)
                r2 = p_title.add_run(src.company)
                r2.font.size = Pt(self.style.body_font_size + 0.5)
                for bullet in exp_adapted.bullets:
                    p = cell.add_paragraph(style=None)
                    p.paragraph_format.left_indent = Cm(0.3)
                    self._write_bullet_runs(
                        p,
                        bullet.text,
                        links_by_id.get(bullet.source_id, []),
                    )

            self._two_col_row(doc, left, _writer)

    def _write_bullet_runs(
        self,
        paragraph,
        text: str,
        links: list[BulletLink],
    ) -> None:
        font_size = Pt(self.style.body_font_size)
        prefix_run = paragraph.add_run("• ")
        prefix_run.font.size = font_size
        for segment in split_bullet_with_links(text, links):
            if segment.url:
                _add_hyperlink_run(paragraph, segment.text, segment.url, font_size)
                _add_hyperlink_run(paragraph, " ↗", segment.url, font_size)
            else:
                run = paragraph.add_run(segment.text)
                run.font.size = font_size

    def _write_education(self, doc: DocxDocument) -> None:
        for deg in education_entries_for_english(self.profile.education):
            start = deg.start_date or str(deg.start_year)
            end = deg.end_date or str(deg.end_year)
            left = f"{start} – {end}"

            def _writer(cell, deg: EducationDisplay = deg):
                p = cell.paragraphs[0]
                r1 = p.add_run(deg.title)
                r1.bold = True
                r1.font.size = Pt(self.style.body_font_size + 0.5)
                if deg.field:
                    r2 = p.add_run(f" – {deg.field}")
                    r2.font.size = Pt(self.style.body_font_size + 0.5)
                p2 = cell.add_paragraph()
                r3 = p2.add_run(deg.institution)
                r3.font.size = Pt(self.style.body_font_size)
                r3.font.color.rgb = RGBColor.from_string(self.style.muted_color_hex)

            self._two_col_row(doc, left, _writer)

    def _write_skills(self, doc: DocxDocument, adapted: AdaptedCV) -> None:
        by_id = {c.id: c for c in self.profile.skills.categories}
        selected = self._selected_skills(adapted)
        if selected:
            for cid, skills in selected.items():
                cat = by_id.get(cid)
                if cat is None or not skills:
                    continue
                p = doc.add_paragraph()
                r1 = p.add_run(f"{cat.name}: ")
                r1.bold = True
                r1.font.size = Pt(self.style.body_font_size)
                r2 = p.add_run(", ".join(skills))
                r2.font.size = Pt(self.style.body_font_size)
            return

        profile_id = (
            adapted.skills_profile_id
            if self.profile.skills.profile_by_id(adapted.skills_profile_id) is not None
            else self._infer_skill_profile_id(adapted)
        )
        merged = self.profile.skills.effective_category_skills(profile_id)
        if merged:
            for cid, skills in merged.items():
                cat = by_id.get(cid)
                if cat is None or not skills:
                    continue
                p = doc.add_paragraph()
                r1 = p.add_run(f"{cat.name}: ")
                r1.bold = True
                r1.font.size = Pt(self.style.body_font_size)
                r2 = p.add_run(", ".join(skills))
                r2.font.size = Pt(self.style.body_font_size)
            return

        ordered_ids = [cid for cid in adapted.skills_order if cid in by_id] + [
            cid for cid in by_id if cid not in adapted.skills_order
        ]
        for cid in ordered_ids:
            cat = by_id[cid]
            p = doc.add_paragraph()
            r1 = p.add_run(f"{cat.name}: ")
            r1.bold = True
            r1.font.size = Pt(self.style.body_font_size)
            r2 = p.add_run(", ".join(cat.skills))
            r2.font.size = Pt(self.style.body_font_size)

    def _selected_skills(self, adapted: AdaptedCV) -> dict[str, list[str]]:
        if not adapted.selected_skills:
            return {}
        by_id = {c.id: c for c in self.profile.skills.categories}
        canonical = {skill.lower(): skill for skill in self.profile.skills.allowed_skills}
        ordered_ids = [cid for cid in adapted.skills_order if cid in by_id] + [
            block.category_id
            for block in adapted.selected_skills
            if block.category_id in by_id and block.category_id not in adapted.skills_order
        ]
        selected_by_category = {
            block.category_id: block.skills for block in adapted.selected_skills
        }
        selected: dict[str, list[str]] = {}
        for cid in ordered_ids:
            skills: list[str] = []
            seen: set[str] = set()
            for skill in selected_by_category.get(cid, []):
                canonical_skill = canonical.get(skill.lower())
                if canonical_skill and canonical_skill not in seen:
                    skills.append(canonical_skill)
                    seen.add(canonical_skill)
            if skills:
                selected[cid] = skills
        return selected

    def _infer_skill_profile_id(self, adapted: AdaptedCV) -> str | None:
        context_bits = [
            adapted.cv_title,
            adapted.professional_summary,
            " ".join(adapted.selected_project_ids),
            " ".join(
                bullet.text
                for exp in adapted.selected_experiences
                for bullet in exp.bullets
            ),
        ]
        return infer_skill_profile_id(self.profile, " ".join(context_bits))

    def _write_projects(self, doc: DocxDocument, adapted: AdaptedCV) -> None:
        by_id = {p.id: p for p in self.profile.projects}
        for pid in adapted.selected_project_ids:
            proj = by_id.get(pid)
            if proj is None:
                continue
            self._write_project_line(doc, proj)

    def _write_project_line(self, doc: DocxDocument, proj: Project) -> None:
        p = doc.add_paragraph()
        bullet_run = p.add_run("• ")
        bullet_run.font.size = Pt(self.style.body_font_size)
        name = p.add_run(proj.name)
        name.bold = True
        name.font.size = Pt(self.style.body_font_size)
        desc = p.add_run(f"  {proj.description}")
        desc.font.size = Pt(self.style.body_font_size)

    def _write_languages(self, doc: DocxDocument) -> None:
        p = doc.add_paragraph()
        bits = []
        for lang in self.profile.languages:
            bits.append(f"{lang.name} — {lang.level}")
        r = p.add_run(" | ".join(bits))
        r.font.size = Pt(self.style.body_font_size)

    def _write_certificates(self, doc: DocxDocument) -> None:
        for cert in self.profile.certificates:
            p = doc.add_paragraph()
            r1 = p.add_run(cert.name)
            r1.bold = True
            r1.font.size = Pt(self.style.body_font_size)
            if cert.url:
                _add_hyperlink_run(p, " ↗", str(cert.url), Pt(self.style.body_font_size))
            r2 = p.add_run(f" — {cert.issuer}")
            r2.font.size = Pt(self.style.body_font_size)
            if cert.description:
                r3 = p.add_run(f" — {cert.description}")
                r3.font.size = Pt(self.style.body_font_size)
                r3.font.color.rgb = RGBColor.from_string(self.style.muted_color_hex)
